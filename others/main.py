import docx


def get_data(re_content, file_path):

    # 获取.docx
    docx_data_all = docx.Document(file_path)

    # 定义变量
    data_dict = {}
    title_data = []
    content_data = []
    data_to_write = []
    doc = docx.Document()

    # 处理原始文本部分
    for para in docx_data_all.paragraphs:
        if para.style.name == 'Heading 3':
            title_data.append(para.text)
            content_data = []
        else:
            if para.text != '':
                content_data.append(para.text)
                data_dict[title_data[-1]] = content_data

    # 获取检索结果部分
    for re in re_content.split(','):
        for item in data_dict.items():
            if re in ''.join(item[1]):
                if item not in data_to_write:
                    data_to_write.append(item)

    # 写入word部分
    for content in data_to_write:
        doc.add_paragraph(content[0], style='Heading 3')
        doc.add_paragraph(''.join(content[1]), style='Normal')
    doc.save(re_content + '_find_all_res.docx')


if __name__ == '__main__':
    get_data(re_content=input('请输入你要查找的内容(任意字符，多词请用英文逗号分割)'), file_path=input('请输入.docx的绝对地址(示例:/home/ivan/Desktop/pythonProject/others/bohai_yifan.docx)'))