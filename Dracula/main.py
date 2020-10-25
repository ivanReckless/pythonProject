import requests
import os
import urllib
from bs4 import BeautifulSoup
from tqdm import tqdm
import json
import re
import docx


def get_title_para():
    data_dict = {}
    data_key_list = []
    data_value = []
    content = docx.Document('/home/ivan/Desktop/pythonProject/Dracula/1.docx')
    for item in content.paragraphs:
        # print(item.style.name)
        if item.style.name == 'Heading 3':
            data_key = item.text
            if data_key not in data_key_list:
                data_key_list.append(data_key)
            else:
                print(data_key)
            data_value = []
        elif item.style.name == 'Normal':
            data_value.append(item.text + 'p')
            data_dict[data_key_list[-1]] = data_value
    json.dump(data_dict, open('data.test.json', 'w', encoding='utf8'))


def get_data():
    url = 'https://seg.shenshen.wiki/api.php'
    # docx_new = docx.Document('/home/ivan/Desktop/pythonProject/Dracula/test_data.docx')
    # for content in tqdm(json.load(open('data.test.json')).items()):
    #     if len(''.join(content[1])) > 400:
    #         with open('error.data.file', 'a', encoding='utf8') as fp:
    #             fp.write(content[0] + '_错误_' + ''.join(content[1]) + '\n')
    #         docx_new.add_paragraph(content[0] + '\n', style='Heading 3')
    #     elif content[0] in open('reserved_file.file').read():
    #         pass
    #     else:
    #         send_dict = {"text": ''.join(content[1])}
    #         send_str = json.dumps(send_dict)
    #         data = {
    #             "text": send_str,
    #             "seg": "biaodian"
    #         }
    #         headers = {
    #             "Host": "seg.shenshen.wiki",
    #             "User-Agent": "Mozilla/5.0 (X11; Linux x86_64; rv:81.0) Gecko/20100101 Firefox/81.0",
    #             "Accept": "*/*",
    #             "Accept-Language": "zh-CN,zh;q=0.8,zh-TW;q=0.7,zh-HK;q=0.5,en-US;q=0.3,en;q=0.2",
    #             "Accept-Encoding": "utf8",
    #             "Content-type": "application/x-www-form-urlencoded",
    #             "Content-Length": "30",
    #             "Origin": "https://seg.shenshen.wiki",
    #             "Connection": "keep-alive",
    #             "Referer": "https://seg.shenshen.wiki/",
    #             "Cookie": "__cfduid=dc99e00f469a1193f91598000c489b9171602484875",
    #             "TE": "Trailers",
    #             "Pragma": "no-cache",
    #             "Cache-Control": "no-cache"
    #         }
    #         r = requests.post(url, data=data, headers=headers).content.decode()
    #         re_content = r'\<.*?\>'
    #         res = re.sub(re_content, '', r)
    #         docx_new.add_paragraph(content[0] + '\n', style='Heading 3')
    #         docx_new.add_paragraph(''.join(res) + '\n', style='Normal')
    #         docx_new.save('/home/ivan/Desktop/pythonProject/Dracula/test_data.docx')
    #         with open('reserved_file.file', 'a', encoding='utf8') as fp:
    #             fp.write(content[0] + '\n')


def get_lost_data():
    doc = docx.Document('/home/ivan/Desktop/pythonProject/Dracula/test_data.docx')
    for content in doc.paragraphs:
        print(content.style.name)


def view_data():
    print(json.load(open('data.test.json')))
    # doc = docx.Document()
    # for para in docx.Document('/home/ivan/Desktop/pythonProject/Dracula/1.docx').paragraphs:
    #     if para.style.name == 'Normal':
    #         doc.add_paragraph(para.text + 'p', style='Normal')
    #     else:
    #         doc.add_paragraph(para.text, style='Heading 3')
    # doc.save('/home/ivan/Desktop/pythonProject/Dracula/2.docx')


def test_main():
    url = 'https://seg.shenshen.wiki/api.php'
    doc = docx.Document('/home/ivan/Desktop/pythonProject/Dracula/test_data_new.docx')
    for i in tqdm(json.load(open('data.test.json')).items()):
        if i[0] in open('reserved_file.file').read():
            continue
        else:
            max_len = 400
            item = ''.join(i[1])
            data_list = [item[idx*max_len:(idx+1)*max_len] for idx in range((len(item) // max_len) + 1)]
            final_data_list = []
            for content in data_list:
                post_dict = {"text": content}
                send_dict = json.dumps(post_dict)
                data = {
                    "text": send_dict,
                    "seg": "biaodian"
                }
                headers = {
                    "Host": "seg.shenshen.wiki",
                    "User-Agent": "Mozilla/5.0 (X11; Linux x86_64; rv:81.0) Gecko/20100101 Firefox/81.0",
                    "Accept": "*/*",
                    "Accept-Language": "zh-CN,zh;q=0.8,zh-TW;q=0.7,zh-HK;q=0.5,en-US;q=0.3,en;q=0.2",
                    "Accept-Encoding": "utf8",
                    "Content-type": "application/x-www-form-urlencoded",
                    "Content-Length": "30",
                    "Origin": "https://seg.shenshen.wiki",
                    "Connection": "keep-alive",
                    "Referer": "https://seg.shenshen.wiki/",
                    "Cookie": "__cfduid=dc99e00f469a1193f91598000c489b9171602484875",
                    "TE": "Trailers",
                    "Pragma": "no-cache",
                    "Cache-Control": "no-cache"
                }
                r = requests.post(url, data=data, headers=headers).content.decode()
                re_content = r'\<.*?\>'
                res = re.sub(re_content, '', r)
                final_data_list.append(res)
            doc.add_paragraph(i[0] + '\n', style='Heading 3')
            doc.add_paragraph(''.join(final_data_list).replace('p', '\n') + '\n', style='Normal')
            doc.save('/home/ivan/Desktop/pythonProject/Dracula/test_data_new.docx')
            with open('reserved_file.file', 'a', encoding='utf8') as fp:
                fp.write(i[0] + '\n')


if __name__ == '__main__':
    # get_data()
    test_main()
    # get_title_para()
    # view_data()
    # get_lost_data()